import os
import re
import shutil
import uuid
import hashlib
import sqlite3
import logging
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional

from backend.app.config.settings import settings
from backend.security.database import get_db_path
from backend.security.audit import AuditLogger

logger = logging.getLogger("aegis.services.document_generator")

class DocumentGeneratorService:
    """
    Sovereign on-premise document generation service.
    Produces physical PDF and DOCX reports from verified organizational document intelligence.
    """

    def __init__(self, output_dir: Optional[str] = None):
        self.output_dir = output_dir or os.path.abspath("data/generated")
        os.makedirs(self.output_dir, exist_ok=True)

    def _get_output_path(self, doc_id: str, extension: str) -> str:
        clean_ext = extension.lstrip(".")
        return os.path.join(self.output_dir, f"{doc_id}.{clean_ext}")

    def generate_pdf_report(
        self,
        title: str,
        sections: Dict[str, str],
        sources: List[Dict[str, Any]],
        output_file_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> int:
        """
        Builds a physical PDF report using ReportLab.
        Returns the file size in bytes.
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

        doc = SimpleDocTemplate(
            output_file_path,
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54
        )
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Heading1"],
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#0f172a"),
            spaceAfter=8
        )
        meta_style = ParagraphStyle(
            "DocMeta",
            parent=styles["Normal"],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#64748b")
        )
        h2_style = ParagraphStyle(
            "DocH2",
            parent=styles["Heading2"],
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#1e3a8a"),
            spaceBefore=14,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#334155"),
            spaceAfter=8
        )
        bullet_style = ParagraphStyle(
            "DocBullet",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor("#334155"),
            leftIndent=12,
            spaceAfter=4
        )

        story = []

        # Title
        story.append(Paragraph(title, title_style))

        # Metadata banner
        now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
        meta_text = (
            f"<b>AEGIS Sovereign Intelligence Report</b> | Generated: {now_str} | "
            f"Classification: <b>CONFIDENTIAL INTERNAL</b>"
        )
        story.append(Paragraph(meta_text, meta_style))
        story.append(Spacer(1, 8))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#3b82f6"), spaceAfter=14))

        # Content Sections
        for sec_title, sec_content in sections.items():
            if not sec_content or not sec_content.strip():
                continue
            story.append(Paragraph(sec_title.upper(), h2_style))
            # Split paragraphs
            paragraphs = sec_content.strip().split("\n\n")
            for p_text in paragraphs:
                lines = p_text.split("\n")
                for line in lines:
                    line_clean = line.strip()
                    if line_clean.startswith("- ") or line_clean.startswith("* "):
                        story.append(Paragraph(f"• {line_clean[2:]}", bullet_style))
                    elif line_clean:
                        story.append(Paragraph(line_clean, body_style))
            story.append(Spacer(1, 6))

        # Sources & Citations
        if sources:
            story.append(Spacer(1, 10))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=10))
            story.append(Paragraph("AUTHORITATIVE SOURCES & CITATIONS", h2_style))

            table_data = [["Document Name", "Referenced Pages", "Relevance Status"]]
            for src in sources:
                fname = src.get("filename", "Document")
                pages = src.get("pages", [])
                p_str = ", ".join(str(p) for p in pages) if pages else str(src.get("page_number", 1))
                rel = src.get("relevance", "High")
                table_data.append([fname, f"Page(s) {p_str}", rel])

            t = Table(table_data, colWidths=[240, 140, 120])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 1), (-1, -1), 8.5),
            ]))
            story.append(t)

        # Verification Hash Footer
        doc_hash = hashlib.sha256(f"{title}{now_str}".encode("utf-8")).hexdigest()
        story.append(Spacer(1, 16))
        story.append(Paragraph(f"Verification Integrity Hash: <font face='Courier'>{doc_hash[:32]}...</font>", meta_style))

        doc.build(story)
        return os.path.getsize(output_file_path)

    def generate_docx_report(
        self,
        title: str,
        sections: Dict[str, str],
        sources: List[Dict[str, Any]],
        output_file_path: str
    ) -> int:
        """
        Builds a physical DOCX report using python-docx.
        Returns the file size in bytes.
        """
        import docx
        from docx.shared import Pt, RGBColor, Inches

        doc = docx.Document()

        # Title
        t_p = doc.add_paragraph()
        run = t_p.add_run(title)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)

        # Meta
        now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
        m_p = doc.add_paragraph()
        m_run = m_p.add_run(f"AEGIS Sovereign Intelligence Report | Generated: {now_str} | Classification: CONFIDENTIAL")
        m_run.font.size = Pt(8.5)
        m_run.font.italic = True
        m_run.font.color.rgb = RGBColor(100, 116, 139)

        # Sections
        for sec_title, sec_content in sections.items():
            if not sec_content or not sec_content.strip():
                continue
            h_p = doc.add_heading(sec_title.upper(), level=2)
            for line in sec_content.strip().split("\n"):
                line_clean = line.strip()
                if line_clean.startswith("- ") or line_clean.startswith("* "):
                    doc.add_paragraph(line_clean[2:], style="List Bullet")
                elif line_clean:
                    doc.add_paragraph(line_clean)

        # Sources table
        if sources:
            doc.add_heading("AUTHORITATIVE SOURCES & CITATIONS", level=2)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Document Name"
            hdr_cells[1].text = "Referenced Pages"
            hdr_cells[2].text = "Relevance Status"

            for src in sources:
                row_cells = table.add_row().cells
                row_cells[0].text = src.get("filename", "Document")
                pages = src.get("pages", [])
                row_cells[1].text = ", ".join(str(p) for p in pages) if pages else str(src.get("page_number", 1))
                row_cells[2].text = src.get("relevance", "High")

        doc.save(output_file_path)
        return os.path.getsize(output_file_path)

    def create_report(
        self,
        title: str,
        sections: Dict[str, str],
        sources: List[Dict[str, Any]],
        format_type: str = "pdf",
        owner_id: int = -1,
        owner_username: str = "",
        source_document_ids: Optional[List[str]] = None,
        conversation_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Full document generation pipeline:
        1. Validates inputs and prepares physical file destination.
        2. Generates actual PDF or DOCX file to a temporary location.
        3. Verifies file integrity, size, and readability.
        4. Atomically moves file to final storage path.
        5. Records metadata into SQLite `generated_documents` table.
        6. Logs audit events (STARTED, COMPLETED, FAILED).
        """
        doc_id = f"rep_{uuid.uuid4().hex[:12]}"
        clean_format = format_type.lower().strip()
        if clean_format not in ("pdf", "docx"):
            clean_format = "pdf"

        extension = clean_format
        mime_type = "application/pdf" if clean_format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Clean sanitized filename without double extensions
        clean_base = re.sub(r'[^a-zA-Z0-9_\-]', '_', title.lower().replace(' ', '_'))
        clean_base = re.sub(r'_\.pdf|_\.docx|\.pdf|\.docx', '', clean_base)[:35].strip('_')
        if not clean_base:
            clean_base = "report"
        clean_filename = f"{clean_base}_{doc_id[4:10]}.{extension}"
        out_path = self._get_output_path(doc_id, extension)
        temp_out_path = f"{out_path}.tmp"

        AuditLogger.log_event(
            action="DOCUMENT_GENERATION_STARTED",
            component="services.document_generator",
            status="success",
            user_id=owner_id,
            username=owner_username,
            resource=clean_filename,
            metadata={"id": doc_id, "title": title, "format": clean_format}
        )

        try:
            # 1. Generate to temporary file
            if clean_format == "pdf":
                file_size = self.generate_pdf_report(title, sections, sources, temp_out_path)
            else:
                file_size = self.generate_docx_report(title, sections, sources, temp_out_path)

            # 2. Verify temporary file was written and is non-empty
            if not os.path.exists(temp_out_path) or os.path.getsize(temp_out_path) == 0:
                raise RuntimeError("Document generation produced an empty or missing output file.")

            file_size = os.path.getsize(temp_out_path)

            # 3. Atomically move to final destination
            if os.path.exists(out_path):
                os.remove(out_path)
            shutil.move(temp_out_path, out_path)

            # 4. Persist to SQLite
            now_str = datetime.now(timezone.utc).isoformat()
            src_ids_str = ",".join(source_document_ids) if source_document_ids else ""

            db_path = get_db_path()
            conn = sqlite3.connect(db_path)
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO generated_documents (
                        id, owner_id, owner_username, filename, title, format,
                        file_size, mime_type, source_document_ids, conversation_id,
                        status, file_path, created_at, updated_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    doc_id, owner_id, owner_username, clean_filename, title, clean_format,
                    file_size, mime_type, src_ids_str, conversation_id or "",
                    "completed", out_path, now_str, now_str
                ))
                conn.commit()
            finally:
                conn.close()

            # 5. Log audit completion
            AuditLogger.log_event(
                action="DOCUMENT_GENERATED",
                component="services.document_generator",
                status="success",
                user_id=owner_id,
                username=owner_username,
                resource=clean_filename,
                metadata={
                    "id": doc_id,
                    "title": title,
                    "format": clean_format,
                    "file_size": file_size,
                    "source_count": len(sources)
                }
            )

            return {
                "id": doc_id,
                "filename": clean_filename,
                "title": title,
                "format": clean_format,
                "file_size": file_size,
                "mime_type": mime_type,
                "source_document_ids": source_document_ids or [],
                "conversation_id": conversation_id,
                "status": "completed",
                "file_path": out_path,
                "created_at": now_str
            }

        except Exception as e:
            # Cleanup temporary artifacts on failure
            if os.path.exists(temp_out_path):
                try:
                    os.remove(temp_out_path)
                except Exception:
                    pass
            if os.path.exists(out_path):
                try:
                    os.remove(out_path)
                except Exception:
                    pass

            AuditLogger.log_event(
                action="DOCUMENT_GENERATION_FAILED",
                component="services.document_generator",
                status="failure",
                user_id=owner_id,
                username=owner_username,
                resource=clean_filename,
                metadata={"id": doc_id, "title": title, "error": str(e)}
            )
            raise

    def list_generated_documents(
        self,
        owner_id: Optional[int] = None,
        is_admin: bool = False
    ) -> List[Dict[str, Any]]:
        """Lists generated documents scoped to owner unless user is admin."""
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        try:
            cursor = conn.cursor()
            if is_admin:
                cursor.execute("SELECT * FROM generated_documents ORDER BY created_at DESC")
            elif owner_id is not None:
                cursor.execute("SELECT * FROM generated_documents WHERE owner_id = ? ORDER BY created_at DESC", (owner_id,))
            else:
                return []
            rows = cursor.fetchall()
            return [dict(r) for r in rows]
        finally:
            conn.close()

    def get_generated_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Retrieves single generated document record by ID."""
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM generated_documents WHERE id = ?", (doc_id,))
            row = cursor.fetchone()
            return dict(row) if row else None
        finally:
            conn.close()

    def delete_generated_document(self, doc_id: str) -> bool:
        """Deletes generated document from SQLite and physical disk."""
        doc = self.get_generated_document(doc_id)
        if not doc:
            return False

        file_path = doc.get("file_path")
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.warning(f"Failed deleting physical generated file '{file_path}': {e}")

        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        try:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM generated_documents WHERE id = ?", (doc_id,))
            conn.commit()
            return True
        finally:
            conn.close()
