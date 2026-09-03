import os
import re
import time
import json
import hashlib
import sqlite3
import logging
from typing import List, Dict, Any, Optional
import pypdf
import chromadb
from chromadb.api.types import EmbeddingFunction, Documents, Embeddings
from backend.rag.embeddings import BaseEmbeddingModel
from backend.multimodal.ocr import BaseOCR, LocalPytesseractOCR
from backend.rag.detector import FileDetector, DetectionResult
from backend.rag.extractors import UniversalExtractorRegistry, NormalizedDocument, NormalizedPage, InsufficientTextError
from backend.security.database import get_db_path

logger = logging.getLogger("aegis.rag")

class RagError(Exception):
    """Base exception for RAG pipeline errors."""
    pass

class SafePathViolationError(RagError):
    """Raised when ingestion path violates configured directory boundary parameters."""
    pass

class DuplicateIngestionError(RagError):
    """Raised when trying to ingest a document that is already parsed and stored."""
    pass

class ChromaEmbeddingAdapter(EmbeddingFunction):
    """Adapts Aegis BaseEmbeddingModel to ChromaDB's internal EmbeddingFunction format."""
    
    def __init__(self, model: BaseEmbeddingModel):
        self.model = model

    def name(self) -> str:
        return getattr(self.model, "model_name", "aegis_embedding")

    def __call__(self, input: Documents) -> Embeddings:
        return self.model.embed_documents(input)

class DocumentLoader:
    """Loads and extracts structured text contents from local PDF, DOCX, TXT, MD, and CSV documents."""
    
    @classmethod
    def load(cls, file_path: str) -> List[Dict[str, Any]]:
        """Alias for load_document."""
        return cls.load_document(file_path)

    @staticmethod
    def load_document(file_path: str) -> List[Dict[str, Any]]:
        """Resolves file extension and returns list of page dictionaries with exact page numbers and order."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".txt" or ext == ".log":
            return DocumentLoader.load_txt(file_path)
        elif ext == ".pdf":
            return DocumentLoader.load_pdf(file_path)
        elif ext == ".docx":
            return DocumentLoader.load_docx(file_path)
        elif ext in [".md", ".markdown"]:
            return DocumentLoader.load_md(file_path)
        elif ext == ".csv":
            return DocumentLoader.load_csv(file_path)
        else:
            raise ValueError(f"Unsupported document file extension format: '{ext}'. Supported formats: .pdf, .docx, .txt, .md, .csv")

    @staticmethod
    def load_txt(file_path: str) -> List[Dict[str, Any]]:
        """Reads plain text file inputs with encoding fallbacks and logical pagination."""
        text = ""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    text = f.read()
            except Exception as e:
                raise ValueError(f"Failed to decode text document: {e}")
                
        if not text.strip():
            raise InsufficientTextError("Document is empty or contains no extractable text characters.")
            
        # Segment into ~2500 character logical pages to keep page numbering grounded
        pages = []
        page_size = 2500
        paragraphs = text.split("\n\n")
        current_page_text = []
        current_len = 0
        page_num = 1
        
        for p in paragraphs:
            p_str = p.strip()
            if not p_str:
                continue
            if current_len + len(p_str) > page_size and current_page_text:
                pages.append({"text": "\n\n".join(current_page_text), "page_number": page_num, "section": "General"})
                page_num += 1
                current_page_text = [p_str]
                current_len = len(p_str)
            else:
                current_page_text.append(p_str)
                current_len += len(p_str)
                
        if current_page_text:
            pages.append({"text": "\n\n".join(current_page_text), "page_number": page_num, "section": "General"})
            
        return pages if pages else [{"text": text.strip(), "page_number": 1, "section": "General"}]

    @staticmethod
    def load_pdf(file_path: str) -> List[Dict[str, Any]]:
        """Extracts text pages from native PDF files using pypdf preserving exact 1-indexed page numbers."""
        pages = []
        total_text = ""
        try:
            reader = pypdf.PdfReader(file_path)
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                clean_page_text = page_text.strip()
                total_text += clean_page_text
                pages.append({
                    "text": clean_page_text,
                    "page_number": idx + 1,
                    "section": f"Page {idx + 1}"
                })
        except Exception as e:
            raise ValueError(f"Failed to parse PDF document contents: {e}")

        # Insufficient text check (indicates a scanned image PDF requiring local OCR)
        if len(total_text.strip()) < 15:
            raise InsufficientTextError("PDF contains no extractable text. Scanned document requires OCR.")
            
        return pages

    @staticmethod
    def load_docx(file_path: str) -> List[Dict[str, Any]]:
        """Extracts structured text from Microsoft Word (.docx) documents preserving headings and tables."""
        try:
            import docx
            doc = docx.Document(file_path)
        except ImportError:
            raise ValueError("python-docx library is required to parse .docx documents.")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {e}")

        pages = []
        current_text = []
        current_page = 1
        current_section = "Introduction"
        current_len = 0
        page_size = 2000

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            if p.style and "Heading" in p.style.name:
                current_section = text
            
            if current_len + len(text) > page_size and current_text:
                pages.append({
                    "text": "\n\n".join(current_text),
                    "page_number": current_page,
                    "section": current_section
                })
                current_page += 1
                current_text = [text]
                current_len = len(text)
            else:
                current_text.append(text)
                current_len += len(text)

        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))
            if table_lines:
                tbl_text = "\n".join(table_lines)
                current_text.append(tbl_text)
                current_len += len(tbl_text)

        if current_text:
            pages.append({
                "text": "\n\n".join(current_text),
                "page_number": current_page,
                "section": current_section
            })

        if not pages or not any(p["text"].strip() for p in pages):
            raise InsufficientTextError("DOCX document contains no extractable text.")

        return pages

    @staticmethod
    def load_md(file_path: str) -> List[Dict[str, Any]]:
        """Extracts structured text from Markdown documents preserving headers and logical sections."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()

        if not content.strip():
            raise InsufficientTextError("Markdown document is empty.")

        sections = re.split(r'\n(?=#{1,3}\s+)', content)
        pages = []
        page_num = 1

        for sec in sections:
            sec_clean = sec.strip()
            if not sec_clean:
                continue
            first_line = sec_clean.split("\n", 1)[0].lstrip("#").strip()
            pages.append({
                "text": sec_clean,
                "page_number": page_num,
                "section": first_line or f"Section {page_num}"
            })
            page_num += 1

        return pages if pages else [{"text": content.strip(), "page_number": 1, "section": "General"}]

    @staticmethod
    def load_csv(file_path: str) -> List[Dict[str, Any]]:
        """Extracts tabular records from CSV files."""
        import csv
        pages = []
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                rows = list(reader)
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                reader = csv.reader(f)
                rows = list(reader)

        if not rows:
            raise InsufficientTextError("CSV file is empty.")

        header = " | ".join(rows[0])
        batch_size = 30
        page_num = 1

        for i in range(1, len(rows), batch_size):
            batch = rows[i:i + batch_size]
            lines = [header, "---"]
            for r in batch:
                lines.append(" | ".join(r))
            pages.append({
                "text": "\n".join(lines),
                "page_number": page_num,
                "section": f"Rows {i}-{min(i+batch_size-1, len(rows)-1)}"
            })
            page_num += 1

        return pages if pages else [{"text": header, "page_number": 1, "section": "Header"}]

class RecursiveTextSplitter:
    """
    Intelligently splits text into semantic overlapping chunks respecting paragraph,
    sentence, and clause boundaries to preserve semantic meaning.
    """
    
    def __init__(
        self,
        chunk_size: int = 900,
        chunk_overlap: int = 150,
        separators: Optional[List[str]] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", "? ", "! ", "; ", ", ", " ", ""]

    def split_text(self, text: str) -> List[str]:
        """Performs hierarchical boundary-aware text splitting."""
        text = text.strip()
        if not text:
            return []
        return self._split(text, self.separators)

    def _split(self, text: str, separators: List[str]) -> List[str]:
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []

        if not separators:
            # Fallback to hard character slicing with overlap
            chunks = []
            pos = 0
            while pos < len(text):
                end = min(pos + self.chunk_size, len(text))
                chunks.append(text[pos:end].strip())
                pos += self.chunk_size - self.chunk_overlap
                if pos >= len(text) - self.chunk_overlap:
                    break
            return [c for c in chunks if c]

        separator = separators[0]
        remaining_separators = separators[1:]

        if separator == "":
            splits = list(text)
        else:
            splits = text.split(separator)

        chunks = []
        current_chunk: List[str] = []
        current_length = 0

        for piece in splits:
            if not piece and separator != "":
                continue
            piece_len = len(piece) + (len(separator) if current_chunk else 0)

            if piece_len > self.chunk_size:
                # Sub-split overly large individual pieces
                if current_chunk:
                    joined = separator.join(current_chunk).strip()
                    if joined:
                        chunks.append(joined)
                    current_chunk = []
                    current_length = 0
                sub_chunks = self._split(piece, remaining_separators)
                chunks.extend(sub_chunks)
                continue

            if current_length + piece_len > self.chunk_size:
                if current_chunk:
                    joined = separator.join(current_chunk).strip()
                    if joined:
                        chunks.append(joined)
                    # Compute overlap from end of current_chunk
                    overlap_pieces: List[str] = []
                    overlap_len = 0
                    for p in reversed(current_chunk):
                        if overlap_len + len(p) <= self.chunk_overlap:
                            overlap_pieces.insert(0, p)
                            overlap_len += len(p) + len(separator)
                        else:
                            break
                    current_chunk = overlap_pieces
                    current_length = sum(len(p) for p in current_chunk) + (len(separator) * (len(current_chunk) - 1) if current_chunk else 0)

            current_chunk.append(piece)
            current_length += piece_len

        if current_chunk:
            joined = separator.join(current_chunk).strip()
            if joined:
                chunks.append(joined)

        return [c for c in chunks if len(c.strip()) > 10]


class AegisRagService:
    """Core local RAG service orchestrating vector indexes, embeddings, and similarity retrieval."""
    
    def __init__(
        self,
        embedding_model: BaseEmbeddingModel,
        persist_directory: str = "vectorstore",
        safe_directories: Optional[List[str]] = None,
        ocr_service: Optional[BaseOCR] = None
    ):
        self.embedding_model = embedding_model
        self.persist_directory = os.path.abspath(persist_directory)
        self.ocr_service = ocr_service
        
        # Enforce local workspace path boundaries to prevent traversal leakage
        self.safe_directories = [os.path.abspath(d) for d in (safe_directories or [os.getcwd()])]
        
        # Initialize local persistent SQLite Chroma Client
        self.chroma_client = chromadb.PersistentClient(path=self.persist_directory)
        self.embedding_fn = ChromaEmbeddingAdapter(self.embedding_model)
        
        # Get or create vector collection with explicit Cosine Distance space
        self.collection = self.chroma_client.get_or_create_collection(
            name="aegis_knowledge",
            embedding_function=self.embedding_fn,
            metadata={"hnsw:space": "cosine"}
        )
        self._verify_collection_compatibility()

    def _verify_collection_compatibility(self):
        """Safely detects and purges stale mock vector data if embedding model upgraded to real transformer weights."""
        try:
            if self.collection.count() > 0:
                sample = self.collection.get(limit=1)
                if sample and sample.get("metadatas") and sample["metadatas"]:
                    meta = sample["metadatas"][0]
                    was_mock = meta.get("is_mock", False) or meta.get("embedding_model") == "mock-embedding-384"
                    is_now_real = not getattr(self.embedding_model, "is_mock", False)
                    if was_mock and is_now_real:
                        logger.warning("Mock embeddings detected in ChromaDB collection. Rebuilding collection for real transformer model.")
                        self.chroma_client.delete_collection("aegis_knowledge")
                        self.collection = self.chroma_client.get_or_create_collection(
                            name="aegis_knowledge",
                            embedding_function=self.embedding_fn,
                            metadata={"hnsw:space": "cosine"}
                        )
        except Exception as e:
            logger.warning(f"Collection compatibility verification notice: {e}")

    def _validate_safe_path(self, file_path: str) -> str:
        """Validates that ingestion target lies strictly inside configured directories."""
        abs_path = os.path.abspath(file_path)
        is_safe = any(abs_path.startswith(safe_dir) for safe_dir in self.safe_directories)
        if not is_safe:
            from backend.security.audit import AuditLogger
            AuditLogger.log_event(
                action="DOCUMENT_INGEST",
                component="rag.pipeline",
                status="failure",
                resource=os.path.basename(file_path),
                metadata={"filename": os.path.basename(file_path), "error_category": "safe_path_violation"}
            )
            raise SafePathViolationError(
                f"Path traversal check failed. File path '{file_path}' is outside safe directories."
            )
        return abs_path

    def _sync_sqlite_document(
        self,
        doc_id: str,
        filename: str,
        source_path: str,
        content_hash: str,
        file_size: int,
        chunk_count: int,
        owner_id: Optional[int] = None,
        owner_username: Optional[str] = None,
        status: str = "indexed",
        mime_type: str = "application/octet-stream",
        document_type: str = "document",
        category: str = "document",
        extraction_method: str = "native",
        metadata_json: str = "{}"
    ) -> None:
        """Persists document registry metadata to authoritative SQLite database."""
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        now_str = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
        try:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO documents (
                    id, filename, source_path, content_hash, file_size, mime_type,
                    document_type, category, extraction_method, metadata_json,
                    chunk_count, owner_id, owner_username, status, created_at, updated_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(id) DO UPDATE SET
                    filename = excluded.filename,
                    source_path = excluded.source_path,
                    file_size = excluded.file_size,
                    mime_type = excluded.mime_type,
                    document_type = excluded.document_type,
                    category = excluded.category,
                    extraction_method = excluded.extraction_method,
                    metadata_json = excluded.metadata_json,
                    chunk_count = excluded.chunk_count,
                    owner_id = excluded.owner_id,
                    owner_username = excluded.owner_username,
                    status = excluded.status,
                    updated_at = excluded.updated_at
            """, (
                doc_id, filename, source_path, content_hash, file_size, mime_type,
                document_type, category, extraction_method, metadata_json,
                chunk_count,
                owner_id if owner_id is not None else -1,
                owner_username or "",
                status, now_str, now_str
            ))
            conn.commit()
        finally:
            conn.close()

    def ingest_document(
        self, 
        file_path: str, 
        chunk_size: int = 900, 
        chunk_overlap: int = 150,
        owner_id: Optional[int] = None,
        owner_username: Optional[str] = None,
        original_filename: Optional[str] = None
    ) -> str:
        """
        Universal ingestion entrypoint: Detects file type via binary magic-bytes,
        extracts normalized pages/sheets/slides/code/OCR, chunks text, generates local embeddings,
        and registers the document in ChromaDB and SQLite.
        """
        from backend.security.audit import AuditLogger
        
        # 1. Path Safety & Existence validations
        try:
            abs_path = self._validate_safe_path(file_path)
            if not os.path.exists(abs_path):
                raise FileNotFoundError(f"Document file does not exist on disk: '{file_path}'")
                
            filename = original_filename if original_filename else os.path.basename(abs_path)
            with open(abs_path, "rb") as file_handle:
                content_bytes = file_handle.read()
                content_hash = hashlib.sha256(content_bytes).hexdigest()
            doc_id = content_hash
            file_size = len(content_bytes)

            detection = FileDetector.detect_from_path(abs_path, filename_override=filename)
            if file_size == 0 or detection.file_type == "empty":
                self._sync_sqlite_document(
                    doc_id=doc_id,
                    filename=filename,
                    source_path=abs_path,
                    content_hash=content_hash,
                    file_size=file_size,
                    chunk_count=0,
                    owner_id=owner_id,
                    owner_username=owner_username,
                    status="failed",
                    mime_type="application/octet-stream",
                    document_type="empty",
                    category="unknown",
                    extraction_method="none",
                    metadata_json=json.dumps({"error": "Empty file (0 bytes)"})
                )
                AuditLogger.log_event(
                    action="DOCUMENT_INDEX_FAILED",
                    component="rag.pipeline",
                    status="failure",
                    resource=filename,
                    metadata={"filename": filename, "error_category": "insufficient_text"}
                )
                raise InsufficientTextError("Document is empty or contains no extractable text.")

            if not detection.is_safe:
                raise ValueError(detection.error_reason or f"Blocked dangerous executable or binary payload in '{filename}'.")
            if not detection.is_valid:
                raise ValueError(detection.error_reason or f"Invalid or unsupported file format for '{filename}'.")

            # Record initial PROCESSING state in SQLite
            self._sync_sqlite_document(
                doc_id=doc_id,
                filename=filename,
                source_path=abs_path,
                content_hash=content_hash,
                file_size=file_size,
                chunk_count=0,
                owner_id=owner_id,
                owner_username=owner_username,
                status="processing",
                mime_type=detection.mime_type,
                document_type=detection.file_type,
                category=detection.category,
                extraction_method=detection.extraction_method,
                metadata_json=json.dumps({"detected_format": detection.file_type, "category": detection.category})
            )

            AuditLogger.log_event(
                action="DOCUMENT_INDEX_STARTED",
                component="rag.pipeline",
                status="success",
                resource=filename,
                metadata={"filename": filename, "id": doc_id, "file_size": file_size}
            )
        except Exception as e:
            if not isinstance(e, SafePathViolationError) and not isinstance(e, InsufficientTextError):
                AuditLogger.log_event(
                    action="DOCUMENT_INDEX_FAILED",
                    component="rag.pipeline",
                    status="failure",
                    resource=os.path.basename(file_path),
                    metadata={"filename": os.path.basename(file_path), "error_category": "path_resolution_error"}
                )
            raise

        # 2. Duplicate Ingestion Verification
        existing = self.collection.get(
            where={"content_hash": content_hash},
            limit=1
        )
        if existing and existing.get("ids") and len(existing["ids"]) > 0:
            AuditLogger.log_event(
                action="DOCUMENT_INDEX_FAILED",
                component="rag.pipeline",
                status="failure",
                resource=filename,
                metadata={"filename": filename, "error_category": "duplicate_rejection"}
            )
            raise DuplicateIngestionError(f"Ingestion rejected. Document '{filename}' is already indexed.")

        # 3. Read Page Contents via Universal Extractor Registry
        try:
            norm_doc = UniversalExtractorRegistry.extract_document(
                abs_path,
                filename=filename,
                detection=detection,
                ocr_service=self.ocr_service
            )
            if not norm_doc.pages or not any(p.text.strip() for p in norm_doc.pages):
                raise InsufficientTextError("Document contains no extractable text.")
        except Exception as e:
            err_cat = "insufficient_text" if ("empty" in str(e).lower() or isinstance(e, InsufficientTextError)) else "parse_failure"
            self._sync_sqlite_document(
                doc_id=doc_id,
                filename=filename,
                source_path=abs_path,
                content_hash=content_hash,
                file_size=file_size,
                chunk_count=0,
                owner_id=owner_id,
                owner_username=owner_username,
                status="failed",
                mime_type=detection.mime_type,
                document_type=detection.file_type,
                category=detection.category,
                extraction_method="failed",
                metadata_json=json.dumps({"error": str(e)})
            )
            AuditLogger.log_event(
                action="DOCUMENT_INDEX_FAILED",
                component="rag.pipeline",
                status="failure",
                resource=filename,
                metadata={"filename": filename, "error_category": err_cat}
            )
            if "empty" in str(e).lower() or isinstance(e, InsufficientTextError):
                raise InsufficientTextError("Document contains no extractable text.")
            raise e

        # 4. Perform Chunking Split
        splitter = RecursiveTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        
        ids = []
        documents = []
        metadatas = []
        
        ingest_time = time.time()
        chunk_idx = 0
        
        for page in norm_doc.pages:
            text = page.text.strip()
            if not text:
                continue
            page_num = page.page_number
            section = page.section_title
            
            chunks = splitter.split_text(text)
            for chunk in chunks:
                chunk_id = f"{doc_id}_c{chunk_idx}"
                ids.append(chunk_id)
                documents.append(chunk)
                metadatas.append({
                    "document_id": doc_id,
                    "filename": filename,
                    "document_name": filename,
                    "source_path": abs_path,
                    "source": abs_path,
                    "page_number": page_num,
                    "section": section,
                    "chunk_id": chunk_id,
                    "chunk_index": chunk_idx,
                    "content_hash": content_hash,
                    "category": norm_doc.category,
                    "file_type": norm_doc.file_type,
                    "extraction_method": norm_doc.extraction_method,
                    "ingested_at": ingest_time,
                    "embedding_model": getattr(self.embedding_model, "model_name", "all-MiniLM-L6-v2"),
                    "is_mock": getattr(self.embedding_model, "is_mock", False),
                    "owner_id": owner_id if owner_id is not None else -1,
                    "owner_username": owner_username if owner_username is not None else ""
                })
                chunk_idx += 1
                
        # 5. Insert to vector index in safe batches
        if ids:
            try:
                batch_size = 500
                for b_idx in range(0, len(ids), batch_size):
                    self.collection.add(
                        ids=ids[b_idx:b_idx + batch_size],
                        documents=documents[b_idx:b_idx + batch_size],
                        metadatas=metadatas[b_idx:b_idx + batch_size]
                    )
                
                # Register in authoritative SQLite database
                self._sync_sqlite_document(
                    doc_id=doc_id,
                    filename=filename,
                    source_path=abs_path,
                    content_hash=content_hash,
                    file_size=file_size,
                    chunk_count=chunk_idx,
                    owner_id=owner_id,
                    owner_username=owner_username,
                    status="indexed",
                    mime_type=norm_doc.mime_type,
                    document_type=norm_doc.file_type,
                    category=norm_doc.category,
                    extraction_method=norm_doc.extraction_method,
                    metadata_json=json.dumps(norm_doc.metadata)
                )
                
                logger.info(f"Ingested document '{filename}' successfully (1 logical document, {chunk_idx} chunks indexed).")
                
                AuditLogger.log_event(
                    action="DOCUMENT_INDEX_COMPLETED",
                    component="rag.pipeline",
                    status="success",
                    resource=filename,
                    metadata={"filename": filename, "file_size": file_size, "chunk_count": chunk_idx, "id": doc_id}
                )
                AuditLogger.log_event(
                    action="DOCUMENT_INDEXED",
                    component="rag.pipeline",
                    status="success",
                    resource=filename,
                    metadata={"filename": filename, "file_size": file_size, "chunk_count": chunk_idx, "id": doc_id}
                )
                AuditLogger.log_event(
                    action="DOCUMENT_INGEST",
                    component="rag.pipeline",
                    status="success",
                    resource=filename,
                    metadata={"filename": filename, "file_size": file_size, "chunk_count": chunk_idx}
                )
            except Exception as e:
                self._sync_sqlite_document(
                    doc_id=doc_id,
                    filename=filename,
                    source_path=abs_path,
                    content_hash=content_hash,
                    file_size=file_size,
                    chunk_count=0,
                    owner_id=owner_id,
                    owner_username=owner_username,
                    status="failed",
                    mime_type=norm_doc.mime_type,
                    document_type=norm_doc.file_type,
                    category=norm_doc.category,
                    extraction_method="failed",
                    metadata_json=json.dumps({"error": str(e)})
                )
                AuditLogger.log_event(
                    action="DOCUMENT_INDEX_FAILED",
                    component="rag.pipeline",
                    status="failure",
                    resource=filename,
                    metadata={"filename": filename, "error_category": "vector_insert_failure"}
                )
                AuditLogger.log_event(
                    action="DOCUMENT_INGEST",
                    component="rag.pipeline",
                    status="failure",
                    resource=filename,
                    metadata={"filename": filename, "error_category": "vector_insert_failure"}
                )
                raise e
            
        return doc_id
            
        return doc_id

    def search(
        self, 
        query: str, 
        top_k: int = 3, 
        filter_metadata: Optional[Dict[str, Any]] = None,
        max_distance: Optional[float] = None,
        document_id: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Executes semantic vector similarity search against local ChromaDB with
        cosine distance thresholding, relevance classification, and deduplication.
        """
        from backend.security.audit import AuditLogger
        start_time = time.perf_counter()
        
        if not query or not query.strip():
            AuditLogger.log_event(
                action="RAG_SEARCH",
                component="rag.pipeline",
                status="failure",
                metadata={"error_category": "empty_query"}
            )
            return []

        try:
            cnt = self.collection.count()
            if isinstance(cnt, int) and cnt == 0:
                AuditLogger.log_event(
                    action="RAG_SEARCH",
                    component="rag.pipeline",
                    status="success",
                    metadata={"query_length": len(query), "chunk_count": 0}
                )
                return []
        except Exception:
            pass

        # Build filter query
        where_filter = {}
        if filter_metadata:
            where_filter.update(filter_metadata)
        if document_id:
            where_filter["document_id"] = document_id

        if len(where_filter) > 1:
            where_clause = {"$and": [{k: v} for k, v in where_filter.items()]}
        elif len(where_filter) == 1:
            where_clause = where_filter
        else:
            where_clause = None

        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=top_k,
                where=where_clause
            )
        except Exception as e:
            AuditLogger.log_event(
                action="RAG_SEARCH",
                component="rag.pipeline",
                status="failure",
                metadata={"query_length": len(query), "error_category": "query_failure"}
            )
            logger.error(f"ChromaDB similarity query run failed: {e}")
            return []

        candidates = []
        if results and results.get("ids") and results["ids"] and len(results["ids"][0]) > 0:
            ids = results["ids"][0]
            docs = results["documents"][0]
            metas = results["metadatas"][0]
            raw_distances = results["distances"][0] if results.get("distances") else [0.0] * len(ids)
            
            seen_texts = set()
            for idx in range(len(ids)):
                dist = float(raw_distances[idx])
                text_content = docs[idx].strip()
                
                # Filter out poor matches if max_distance threshold is supplied
                if max_distance is not None and dist > max_distance:
                    continue
                    
                # Deduplicate identical or near-identical text
                text_signature = text_content[:120].lower()
                if text_signature in seen_texts:
                    continue
                seen_texts.add(text_signature)
                
                # Calculate normalized similarity score (0.0 to 1.0)
                similarity = round(max(0.0, min(1.0, 1.0 - dist)), 4)
                
                if similarity >= 0.60:
                    relevance = "High"
                elif similarity >= 0.35:
                    relevance = "Medium"
                else:
                    relevance = "Low"

                candidates.append({
                    "chunk_id": ids[idx],
                    "text": text_content,
                    "metadata": metas[idx],
                    "distance": round(dist, 4),
                    "similarity": similarity,
                    "relevance": relevance
                })
                
                if len(candidates) >= top_k:
                    break

        duration_ms = int((time.perf_counter() - start_time) * 1000)

        AuditLogger.log_event(
            action="RAG_SEARCH",
            component="rag.pipeline",
            status="success",
            duration_ms=duration_ms,
            metadata={"query_length": len(query), "chunk_count": len(candidates), "duration_ms": duration_ms}
        )

        return candidates

    def get_document_chunks(self, doc_id: str) -> List[Dict[str, Any]]:
        """Retrieves all indexed chunks for a given document ID."""
        data = self.collection.get(
            where={"document_id": doc_id},
            include=["documents", "metadatas"]
        )
        if not data or not data.get("ids"):
            return []
            
        chunks = []
        for i in range(len(data["ids"])):
            chunks.append({
                "chunk_id": data["ids"][i],
                "text": data["documents"][i],
                "metadata": data["metadatas"][i]
            })
            
        # Sort by chunk_index
        chunks.sort(key=lambda c: c.get("metadata", {}).get("chunk_index", 0))
        return chunks

    def delete_document(self, doc_id: str, delete_file: bool = False) -> None:
        """Deletes all chunks belonging to a document ID, cleans SQLite registry, and optionally removes disk file."""
        # 1. Look up file path from SQLite before deleting
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        source_path = None
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT source_path FROM documents WHERE id = ?", (doc_id,))
            row = cursor.fetchone()
            if row:
                source_path = row[0]
            cursor.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
            conn.commit()
        finally:
            conn.close()

        # 2. Delete vectors from ChromaDB
        existing = self.collection.get(where={"document_id": doc_id})
        if existing and existing.get("ids") and len(existing["ids"]) > 0:
            self.collection.delete(ids=existing["ids"])

        existing_hash = self.collection.get(where={"content_hash": doc_id})
        if existing_hash and existing_hash.get("ids") and len(existing_hash["ids"]) > 0:
            self.collection.delete(ids=existing_hash["ids"])
            
        # 3. Optionally remove physical file
        if delete_file and source_path and os.path.exists(source_path):
            try:
                os.remove(source_path)
            except Exception as e:
                logger.warning(f"Could not remove physical file at '{source_path}': {e}")
                
        logger.info(f"Successfully deleted document ID '{doc_id}' from vector store and registry.")

    def list_documents(self, owner_id: Optional[int] = None, is_admin: bool = False) -> List[Dict[str, Any]]:
        """
        Retrieves list of all unique logical documents from authoritative SQLite database
        and cross-references with ChromaDB.
        """
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        try:
            cursor = conn.cursor()
            query = """
                SELECT id, filename, source_path, content_hash, file_size, mime_type,
                       document_type, category, extraction_method, metadata_json,
                       chunk_count, owner_id, owner_username, status, created_at, updated_at
                FROM documents
            """
            if is_admin:
                cursor.execute(query + " WHERE status = 'indexed' ORDER BY created_at DESC")
            elif owner_id is not None:
                cursor.execute(query + " WHERE owner_id = ? AND status = 'indexed' ORDER BY created_at DESC", (owner_id,))
            else:
                cursor.execute(query + " WHERE status = 'indexed' ORDER BY created_at DESC")
                
            rows = cursor.fetchall()
            
            docs = []
            for r in rows:
                meta = {}
                if "metadata_json" in r.keys() and r["metadata_json"]:
                    try:
                        meta = json.loads(r["metadata_json"])
                    except Exception:
                        meta = {}

                docs.append({
                    "document_id": r["id"],
                    "id": r["id"],
                    "filename": r["filename"],
                    "source_path": r["source_path"],
                    "content_hash": r["content_hash"],
                    "file_size": r["file_size"],
                    "mime_type": r["mime_type"],
                    "document_type": r["document_type"] if "document_type" in r.keys() else "document",
                    "category": r["category"] if "category" in r.keys() else "document",
                    "extraction_method": r["extraction_method"] if "extraction_method" in r.keys() else "native",
                    "metadata": meta,
                    "chunk_count": r["chunk_count"],
                    "owner_id": r["owner_id"],
                    "owner_username": r["owner_username"],
                    "status": r["status"],
                    "ingested_at": r["created_at"],
                    "uploaded_at": r["created_at"],
                    "created_at": r["created_at"],
                    "updated_at": r["updated_at"],
                    "is_mock": False
                })
            return docs
        finally:
            conn.close()

    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Retrieves a single logical document metadata by ID."""
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        try:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, filename, source_path, content_hash, file_size, mime_type,
                       document_type, category, extraction_method, metadata_json,
                       chunk_count, owner_id, owner_username, status, created_at, updated_at
                FROM documents
                WHERE id = ?
            """, (doc_id,))
            row = cursor.fetchone()
            if not row:
                return None
            r = dict(row)
            meta = {}
            if r.get("metadata_json"):
                try:
                    meta = json.loads(r["metadata_json"])
                except Exception:
                    meta = {}

            return {
                "document_id": r["id"],
                "id": r["id"],
                "filename": r["filename"],
                "source_path": r["source_path"],
                "content_hash": r["content_hash"],
                "file_size": r["file_size"],
                "mime_type": r["mime_type"],
                "document_type": r.get("document_type", "document"),
                "category": r.get("category", "document"),
                "extraction_method": r.get("extraction_method", "native"),
                "metadata": meta,
                "chunk_count": r["chunk_count"],
                "owner_id": r["owner_id"],
                "owner_username": r["owner_username"],
                "status": r["status"],
                "ingested_at": r["created_at"],
                "uploaded_at": r["created_at"]
            }
        finally:
            conn.close()

    def get_document_stats(self, owner_id: Optional[int] = None, is_admin: bool = False) -> Dict[str, Any]:
        """Calculates exact count of documents, chunks, and storage bytes from authoritative SQLite registry."""
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        try:
            cursor = conn.cursor()
            if is_admin:
                cursor.execute("SELECT COUNT(*), COALESCE(SUM(chunk_count), 0), COALESCE(SUM(file_size), 0) FROM documents WHERE status = 'indexed'")
                indexed_row = cursor.fetchone()
                cursor.execute("SELECT COUNT(*) FROM documents WHERE status = 'failed'")
                failed_count = cursor.fetchone()[0]
                cursor.execute("SELECT COUNT(*) FROM documents WHERE status = 'processing'")
                processing_count = cursor.fetchone()[0]
            elif owner_id is not None:
                cursor.execute("SELECT COUNT(*), COALESCE(SUM(chunk_count), 0), COALESCE(SUM(file_size), 0) FROM documents WHERE owner_id = ? AND status = 'indexed'", (owner_id,))
                indexed_row = cursor.fetchone()
                cursor.execute("SELECT COUNT(*) FROM documents WHERE owner_id = ? AND status = 'failed'", (owner_id,))
                failed_count = cursor.fetchone()[0]
                cursor.execute("SELECT COUNT(*) FROM documents WHERE owner_id = ? AND status = 'processing'", (owner_id,))
                processing_count = cursor.fetchone()[0]
            else:
                return {
                    "total_documents": 0,
                    "indexed_documents": 0,
                    "failed_documents": 0,
                    "processing_documents": 0,
                    "total_chunks": 0,
                    "total_file_size": 0
                }
            
            indexed_count = indexed_row[0]
            total_chunks = indexed_row[1]
            total_file_size = indexed_row[2]
            
            return {
                "total_documents": indexed_count + failed_count + processing_count,
                "indexed_documents": indexed_count,
                "failed_documents": failed_count,
                "processing_documents": processing_count,
                "total_chunks": total_chunks,
                "total_file_size": total_file_size
            }
        finally:
            conn.close()

