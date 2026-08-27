import os
import unittest
import shutil
import tempfile
from docx import Document
from openpyxl import load_workbook
import pypdf
from backend.tools.document_generators.generators import (
    DocxGenerator,
    XlsxGenerator,
    PdfGenerator,
    SafePathViolationError
)

class TestAegisDocumentGenerators(unittest.TestCase):
    """Unit tests for offline document deliverable generation pipelines."""
    
    @classmethod
    def setUpClass(cls):
        # Setup temporary directories for outputs testing
        cls.test_outputs_dir = tempfile.mkdtemp()
        
        # Instantiate generators targeting the temp outputs folder
        cls.docx_gen = DocxGenerator(output_base_dir=cls.test_outputs_dir)
        cls.xlsx_gen = XlsxGenerator(output_base_dir=cls.test_outputs_dir)
        cls.pdf_gen = PdfGenerator(output_base_dir=cls.test_outputs_dir)
        
        # Define mock structured contents
        cls.mock_content = [
            {"type": "heading", "text": "Section 1: General Maintenance", "level": 1},
            {"type": "paragraph", "text": "This report details pipeline pressure levels at Mangalore Refinery."},
            {"type": "bullet", "text": "Check all check valves hourly."},
            {"type": "bullet", "text": "Confirm seal integrity daily."},
            {"type": "numbered", "text": "Shut down the pipeline if pressure exceeds 150 PSI."},
            {
                "type": "table",
                "headers": ["Equipment ID", "Status", "Pressure (PSI)"],
                "rows": [
                    ["EQ-101", "Normal", "110"],
                    ["EQ-202", "Under Maintenance", "0"],
                    ["EQ-303", "Warning", "145"]
                ]
            }
        ]
        
        cls.mock_sheets = [
            {
                "name": "Pressure Logs",
                "headers": ["Timestamp", "Reading (PSI)", "Observer"],
                "rows": [
                    ["2026-08-27 10:00", 115.5, "Operator A"],
                    ["2026-08-27 11:00", 120.2, "Operator B"],
                    ["2026-08-27 12:00", 145.0, "Operator A"]
                ]
            },
            {
                "name": "Audit Comments",
                "headers": ["Inspector", "Notes"],
                "rows": [
                    ["Inspector X", "Valves need lubrication"],
                    ["Inspector Y", "Gauge recalibrated successfully"]
                ]
            }
        ]

    @classmethod
    def tearDownClass(cls):
        # Clean up generated test output folder
        if os.path.exists(cls.test_outputs_dir):
            try:
                shutil.rmtree(cls.test_outputs_dir)
            except Exception:
                pass

    def test_docx_generation_and_verification(self):
        """1, 9, 12. Verify DOCX creation and validate structure by reopening."""
        filename = "test_note.docx"
        target_path = self.docx_gen.generate_docx(filename, "Aegis Approval Note", self.mock_content)
        
        # Verify file exists on disk
        self.assertTrue(os.path.exists(target_path))
        self.assertEqual(os.path.basename(target_path), filename)
        
        # Re-open and verify content
        doc = Document(target_path)
        
        # Check title
        self.assertEqual(doc.paragraphs[0].text, "Aegis Approval Note")
        
        # Check heading list
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        self.assertIn("Section 1: General Maintenance", headings)
        
        # Check list items
        bullets = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
        self.assertIn("Check all check valves hourly.", bullets)
        
        # Check table
        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        self.assertEqual(table.rows[0].cells[0].text, "Equipment ID")
        self.assertEqual(table.rows[1].cells[0].text, "EQ-101")

    def test_xlsx_generation_and_verification(self):
        """2, 9, 12. Verify XLSX creation and validate sheets by reopening."""
        filename = "test_sheet.xlsx"
        target_path = self.xlsx_gen.generate_xlsx(filename, self.mock_sheets)
        
        # Verify file exists on disk
        self.assertTrue(os.path.exists(target_path))
        self.assertEqual(os.path.basename(target_path), filename)
        
        # Re-open and verify content
        wb = load_workbook(target_path)
        self.assertIn("Pressure Logs", wb.sheetnames)
        self.assertIn("Audit Comments", wb.sheetnames)
        
        ws = wb["Pressure Logs"]
        # Check headers
        self.assertEqual(ws.cell(row=1, column=1).value, "Timestamp")
        # Check formatting bold
        self.assertTrue(ws.cell(row=1, column=1).font.bold)
        # Check data
        self.assertEqual(ws.cell(row=2, column=2).value, 115.5)

    def test_pdf_generation_and_verification(self):
        """3, 9, 12. Verify PDF creation and validate text content by reopening."""
        filename = "test_doc.pdf"
        target_path = self.pdf_gen.generate_pdf(filename, "Aegis Audit Report", self.mock_content)
        
        # Verify file exists on disk
        self.assertTrue(os.path.exists(target_path))
        self.assertEqual(os.path.basename(target_path), filename)
        
        # Re-open using pypdf and extract text
        reader = pypdf.PdfReader(target_path)
        self.assertGreater(len(reader.pages), 0)
        
        extracted_text = reader.pages[0].extract_text()
        self.assertIn("Aegis Audit Report", extracted_text)
        self.assertIn("Section 1: General Maintenance", extracted_text)
        self.assertIn("EQ-101", extracted_text)

    def test_path_traversal_rejection(self):
        """5. Verify path traversals outside the base outputs directory are blocked."""
        # Absolute path escape attempt
        with self.assertRaises(SafePathViolationError):
            self.docx_gen.generate_docx("../../leak.docx", "Title", self.mock_content)
            
        with self.assertRaises(SafePathViolationError):
            self.xlsx_gen.generate_xlsx("..\\..\\leak.xlsx", self.mock_sheets)

        # Unsafe name mapping basename conversion verification
        target_path = self.docx_gen.generate_docx("doc/../safe_name.docx", "Title", self.mock_content)
        self.assertEqual(os.path.basename(target_path), "safe_name.docx")

    def test_empty_content_handling(self):
        """8. Verify empty inputs compile safely without raising errors."""
        # DOCX empty content list
        target_docx = self.docx_gen.generate_docx("empty.docx", "Empty Doc", [])
        self.assertTrue(os.path.exists(target_docx))
        
        # XLSX empty sheet list
        target_xlsx = self.xlsx_gen.generate_xlsx("empty.xlsx", [])
        self.assertTrue(os.path.exists(target_xlsx))
        
        # PDF empty content list
        target_pdf = self.pdf_gen.generate_pdf("empty.pdf", "Empty Doc", [])
        self.assertTrue(os.path.exists(target_pdf))
